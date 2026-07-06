import os
import re
import json
import hashlib
import datetime
from docx import Document
from dev.diagnostics import get_logger, terminal_alert

logger = get_logger("data_parser.dish_parser")
salert = terminal_alert

OUTPUT_DIR = r"X:\SAFU NOVA\userdata\parsed_data"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def load_raw_document(file_path: str) -> dict:
    logger.info(f"[LOAD_RAW] Loading file: {file_path}")
    try:
        ext = os.path.splitext(file_path)[-1].lower()

        if ext == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            logger.info(f"[LOAD_RAW] Loaded .txt file: {len(content)} characters")

        elif ext == ".docx":
            doc = Document(file_path)
            content = "\n".join(p.text for p in doc.paragraphs)
            logger.info(f"[LOAD_RAW] Loaded .docx file: {len(doc.paragraphs)} paragraphs")

        else:
            raise ValueError("Unsupported file format. Only .txt and .docx are allowed.")

        return {
            "status": "success",
            "file": os.path.basename(file_path),
            "content": content.strip().replace("\r\n", "\n").replace("\r", "\n")
        }

    except Exception as e:
        logger.error(f"[RAW LOAD ERROR] Failed to load raw document: {e}")
        return {"status": "error", "message": str(e)}


def extract_metadata(paragraphs):
    logger.info("[META] Extracting metadata from first 10 lines...")
    metadata = {}
    patterns = {
        "title": re.compile(r"(?i)^title\s*[:\-]\s*(.+)"),
        "chapter": re.compile(r"(?i)^chapter\s*[:\-]\s*(.+)"),
        "pov": re.compile(r"(?i)^pov\s*[:\-]\s*(.+)"),
        "time": re.compile(r"(?i)^time\s*[:\-]\s*(.+)"),
        "date": re.compile(r"(?i)^date\s*[:\-]\s*(.+)")
    }

    for para in paragraphs[:10]:
        line = para.text.strip()
        for key, pattern in patterns.items():
            match = pattern.match(line)
            if match:
                metadata[key] = match.group(1).strip()
                logger.info(f"[META] Found {key}: {metadata[key]}")
    return metadata


def chunk_paragraphs(paragraphs, metadata, doc_hash,
                     min_group_lines=2, max_group_lines=10,
                     max_paragraph_length=1000, overlap_lines=0):
    logger.info(f"[CHUNKING] Starting dynamic chunking with doc_hash: {doc_hash}")
    chunks = []
    chunk_index = 1
    buffer = []
    buffer_line_count = 0

    def _flush_chunk(text):
        nonlocal chunk_index
        text = text.strip()
        if text:
            logger.debug(f"[CHUNK_FLUSH] Chunk #{chunk_index} ({len(text)} chars)")
            preview = text[:100].replace("\n", " ")
            logger.debug(f"[CHUNK_PREVIEW] {preview}...")
            chunks.append({
                "chunk_id": f"{doc_hash}_chunk_{chunk_index:04d}",
                "chunk_number": chunk_index,
                "chunk": f"chunk {chunk_index}:\n{text}",
                "meta": metadata.copy(),
                "doc_hash": doc_hash
            })
            chunk_index += 1

    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if len(text) > max_paragraph_length:
            if buffer:
                _flush_chunk("\n".join(buffer))
                buffer = buffer[-overlap_lines:] if overlap_lines else []
                buffer_line_count = len(buffer)

            start = 0
            while start < len(text):
                end = min(start + max_paragraph_length, len(text))
                _flush_chunk(text[start:end])
                start = end
        else:
            buffer.append(text)
            buffer_line_count += 1

            if buffer_line_count >= max_group_lines:
                if buffer_line_count >= min_group_lines:
                    _flush_chunk("\n".join(buffer))
                    buffer = buffer[-overlap_lines:] if overlap_lines else []
                    buffer_line_count = len(buffer)

    if buffer:
        _flush_chunk("\n".join(buffer))

    logger.info(f"[CHUNKING] Completed with {len(chunks)} total chunks")
    return chunks


def parse_document(file_path: str) -> dict:
    logger.info(f"[PARSING] Parsing document: {file_path}")
    try:
        ext = os.path.splitext(file_path)[-1].lower()
        paragraphs = []

        if ext == ".txt":
            logger.info("[PARSING] Detected .txt format")
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = [line for line in f.readlines() if line.strip()]
                paragraphs = [LineStub(line) for line in lines]
            logger.info(f"[PARSING] Loaded {len(paragraphs)} lines as LineStubs")

        elif ext == ".docx":
            logger.info("[PARSING] Detected .docx format")
            doc = Document(file_path)
            paragraphs = doc.paragraphs
            logger.info(f"[PARSING] Loaded {len(paragraphs)} paragraphs")

        else:
            raise ValueError(f"Unsupported file type: {ext}")

        full_text = "\n".join(p.text.strip() for p in paragraphs if p.text.strip())
        doc_hash = hashlib.sha256(full_text.encode()).hexdigest()[:12]
        logger.info(f"[PARSING] Computed document hash: {doc_hash}")

        metadata = extract_metadata(paragraphs)
        metadata["doc_hash"] = doc_hash

        chunks = chunk_paragraphs(paragraphs, metadata, doc_hash)

        return {
            "filename": os.path.basename(file_path),
            "created_at": datetime.datetime.now().isoformat(),
            "hash_id": doc_hash,
            "metadata": metadata,
            "chunk_count": len(chunks),
            "chunks": chunks
        }

    except Exception as e:
        logger.error(f"[PARSER ERROR] Failed to parse {file_path}: {e}", exc_info=True)
        raise


class LineStub:
    def __init__(self, text):
        self.text = text
        logger.debug(f"[LINESTUB] Created LineStub: {text[:30]}")


def save_parsed_json(parsed_data: dict):
    try:
        base_name = os.path.splitext(parsed_data["filename"])[0]
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{base_name}_{parsed_data['hash_id']}_{timestamp}.json"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(parsed_data, f, indent=2)

        logger.info(f"[SAVE] JSON saved to {output_path}")
        salert(f"[DOC PARSER] Parsed: {output_filename}", level="INFO")

    except Exception as e:
        logger.error(f"[SAVE ERROR] Could not write JSON file: {e}", exc_info=True)


def process_file(file_path: str):
    logger.info(f"[PROCESS] Starting full parse/save for file: {file_path}")
    try:
        parsed = parse_document(file_path)
        save_parsed_json(parsed)
        logger.info(f"[PROCESS] Completed processing of {file_path}")
    except Exception as e:
        salert(f"[DOC PARSER ERROR] {e}", level="ERROR")
        logger.error(f"[PROCESS FAIL] {file_path}: {e}", exc_info=True)
